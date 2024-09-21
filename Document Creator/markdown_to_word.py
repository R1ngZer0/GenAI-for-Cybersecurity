from docx import Document
from markdown2 import markdown

# Function to convert markdown to Word
def convert(detailed_info):
    # Create a new Document
    doc = Document()
    
    # Split the markdown text into lines
    lines = detailed_info.split("\n")

    # Iterate through each line
    for line in lines:
        # Check if the line is a heading
        if line.startswith("#"):
            # Get the heading level
            heading_level = line.count("#")

            # Get the heading text
            heading_text = line.strip("# ")

            # Add the heading to the Word document
            doc.add_heading(heading_text, level=heading_level)

        # Check if the line is a list item
        elif line.startswith("- "):
            # Add the list item to the Word document
            doc.add_paragraph(line.strip("- "), style="List Bullet")

        # Check if the line is a numbered list item
        elif line.startswith("1. "):
            # Add the numbered list item to the Word document
            doc.add_paragraph(line.strip("1. "), style="List Number")

        # Check if the line contains bold or italic text
        elif "**" in line or "*" in line:
            # Split the line into parts for bold and italic handling
            bold_parts = line.split("**")
            paragraph = doc.add_paragraph()

            for i, part in enumerate(bold_parts):
                if i % 2 == 1:
                    run = paragraph.add_run(part)
                    run.bold = True
                else:
                    italic_parts = part.split("*")
                    for j, subpart in enumerate(italic_parts):
                        if j % 2 == 1:
                            run = paragraph.add_run(subpart)
                            run.italic = True
                        else:
                            paragraph.add_run(subpart)

        else:
            # Add the regular text to the Word document
            doc.add_paragraph(line)
