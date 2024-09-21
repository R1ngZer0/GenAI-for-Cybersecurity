import os
import openai
from openai import OpenAI
import docx
# from markdown import markdown
from markdown2 import markdown
from tqdm import tqdm
import json
import sys
import datetime
# from docx import Document
# from docx.shared import Pt


if len(sys.argv) < 2:
    print("Usage: python irPlan.py <file_path>")
    sys.exit(1)

# get the OpenAI API key from environment variable
openai.api_key = os.getenv('OPENAI_API_KEY')
file_path = sys.argv[1]

# Open irPlanOutline.txt, read the contents, split the contents by '--' and store the sections in a list
with open('irPlanOutline.txt', 'r') as file:
    outline = file.read()
    sections = outline.split('---')

# Print each section
# for i, section in enumerate(sections, start=1):
#     print(section.strip())

conversation_history = []

client = OpenAI(
    base_url = 'http://127.0.0.1:11434/v1',
    api_key = 'ollama',
)

# load a JSON file
def load_json(file_path):
    try:
        with open(file_path, 'r') as file:
            data = json.load(file)
        return data
    except json.JSONDecodeError:
        print("Error: Invalid file path. No JSON file found.")
        sys.exit(1)

try:
    # load the JSON file
    data = load_json(file_path)
    company = data["companyName"]
    company_description = data["companyInfo"]
    industry = data["industry"]
    network_details = data["networkEnvironment"]
    it_staff = data["itStaff"]
    ot_staff = data["otStaff"]
    cybersec_staff = data["cybersecurityStaff"]
    leadership_staff = data["leadershipStaff"]
except KeyError as e:
    print(f"Error: Missing expected key {e} in JSON file.")
    sys.exit(1)

# prepare Word document
doc = docx.Document()
html_text = ""

system_prompt = f"You are a cybersecurity professional specializing in incident response (IR) with more than 25 years of experience. Your job is to create a detailed incident response plan, which includes a ransomware playbook at the end, for my company which has the following details: Company Name: {company}, Company Description: {company_description}, Industry: {industry}, IT Staff: {it_staff}, OT Staff: {ot_staff}, Cybersecurity Staff: {cybersec_staff}, Leadership Staff: {leadership_staff}Network Details: {network_details}. The plan should be written in a professional and formal tone, as if it were a real incident response plan for a real organization. The plan should be detailed and comprehensive, covering all aspects of incident response. In addition to being informational, it should be specific to the organization details provided. It should not just be an instructional document educating the reader on IR, but also information as to how the organization plans to deal with and respond to incidents, specifically, based on the given information about the company. Be sure to include an organization 'call list' in the appropriate section (provide fictional details if you don't have that information provided to you). The playbook should be concise with actionable items, details, and instructions that are specific to the details provided (provide as much technical specifics and step-by-step instructions as possible). It doesn't need to be educational since that part is covered in the main parts of the plan. The plan will be created according to the following outline: {outline}"

conversation_history.append({"role": "system", "content": system_prompt})

print(f"\nGenerating IR Plan...\n")

# Initialize the progress bar
progress_bar = tqdm(total=len(sections), desc="Generating Sections", bar_format="{l_bar}{bar}| {n_fmt}/{total_fmt} [{elapsed}] ")

# Process each section in the outline
for section in sections:
    section = section.strip()
    if not section:
        continue

    initial_instruction = f"Write the narrative, context, and details for the following section (and only this section): {section}. Use as much detail and explanation as possible. In addition to being descriptive, it should be specific to the organization in how it will deal with and respond to threats, for that section. Do not write anything that should go in another section of the plan. You should base the details on the information given, but create fictitious details including organizational names, and other details as needed for any missing information. Use markdown language to denote the proper headings, lists, formatting, etc. Do no mention that this is a fictional plan or that the details are not real. The output should be detailed and comprehensive for this section. Once complete, do not ask for feedback or make any changes to the output. Do not add any followup explanation or summary of the output."

    conversation_history.append({"role": "user", "content": initial_instruction})

    # prepare prompt for detailed info
    messages = conversation_history
        
    try:
        response = client.chat.completions.create(
            model="llama3",
            messages=messages,
            temperature=0.5,
        )

        # get detailed info
        detailed_info = response.choices[0].message.content.strip() 

        conversation_history.append({"role": "assistant", "content": detailed_info})

        # Add to Word document
        doc.add_paragraph(detailed_info)
        doc.add_paragraph("\n")  # add extra line break for readability

        # Convert markdown to HTML and add to the html_text string
        html_text += markdown(detailed_info)

        # Update the progress bar
        progress_bar.update(1)

    except Exception as e:
        print("An error occurred while connecting to the OpenAI API:", e)
        exit(1)

# Close the progress bar
progress_bar.close()

# generate date for the filename
date = datetime.datetime.now().strftime("%Y-%m-%d-%H-%M")

# save Word document
print("Saving IR Plan in MS Word Format...")
try:
    doc.save(f"Incident_Response_Plan_{date}.docx")
except PermissionError as e:
    print(f"Error: {e}")
    sys.exit(1)

# save HTML document
print("Saving IR Plan in HTML Format...")
try:
    with open(f"Incident_Response_Plan_{date}.html", 'w') as f:
        f.write(html_text)
except PermissionError as e:
    print(f"Error: {e}")
    sys.exit(1)

print("\nDone.")